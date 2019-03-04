from docx import Document
import os.path
import comtypes.client

def createPDF(pathDir,pathDir2):
    wdFormatPDF = 17

    in_file = os.path.abspath(pathDir)
    out_file = os.path.abspath(pathDir2)

    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(in_file)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()



def createDoc(list,pathDir):

    document = Document()

    if list[0]=="University Exam":
        document.add_heading(list[0], 0)
        document.add_heading('Part A', level=2)

        for i in range(len(list[1])):
            document.add_paragraph(list[1][i][0],style='List Number')

            for j in range(1,len(list[1][i])):
                document.add_paragraph((chr)(97+j-1)+") " + str(list[1][i][j]), style="List 3")

        document.add_heading('Part B', level=2)
        for i in range(len(list[2])):
            document.add_paragraph(list[2][i], style='List Number')

        document.add_heading('Part C', level=2)
        for i in range(len(list[3])):
            document.add_paragraph(list[3][i][0], style='List Number')
            p=document.add_paragraph('OR')
            p.paragraph_format.alignment = 1
            document.add_paragraph(list[3][i][1], style='List 2')
    else:
        document.add_heading(list[0], 0)
        document.add_heading('Part A', level=2)
        for i in range(len(list[1])):
            document.add_paragraph(list[1][i], style='List Number')

        document.add_heading('Part B', level=2)
        for i in range(len(list[2])):
            document.add_paragraph(list[2][i][0], style='List Number')
            p = document.add_paragraph('OR')
            p.paragraph_format.alignment = 1
            document.add_paragraph(list[2][i][1], style='List 2')

    pathDir2 = pathDir + '/Questions.pdf'
    pathDir = pathDir + '/Questions.docx'
    document.save(pathDir)
    createPDF(pathDir,pathDir2)


    if os.path.isfile(pathDir2) == True:
        return True
    else:
        return False



